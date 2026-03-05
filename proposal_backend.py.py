from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

app = Flask(__name__)
CORS(app)

def create_proposal_doc(data):
    """Generate a complete proposal Word document from scratch"""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # COVER PAGE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Proposal for the\n")
    p.add_run("Buyer Plant Machine Trains\n").bold = True
    p.add_run("\nSETPOINT™ Machinery Protection Systems\nAnd Condition Monitoring Software\n")
    
    doc.add_paragraph()
    
    # Buyer Contact
    doc.add_paragraph("Buyer Contact:").runs[0].bold = True
    doc.add_paragraph(data.get("buyerContactName", ""))
    doc.add_paragraph(data.get("buyerContactTitle", ""))
    doc.add_paragraph(data.get("buyerCompanyName", ""))
    doc.add_paragraph(data.get("buyerCompanyAddress", ""))
    
    doc.add_paragraph()
    
    # Sales Contact
    doc.add_paragraph("Brüel & Kjær Vibro Contact:").runs[0].bold = True
    doc.add_paragraph(data.get("salesName", ""))
    doc.add_paragraph(data.get("salesTitle", ""))
    doc.add_paragraph("www.bkvibro.com")
    
    doc.add_page_break()
    
    # PAGE 2: PRICING
    doc.add_heading("2. Pricing", level=1)
    
    pricing_table = doc.add_table(rows=1, cols=2)
    pricing_table.style = 'Light Grid Accent 1'
    hdr_cells = pricing_table.rows[0].cells
    hdr_cells[0].text = "Scope of Supply Description"
    hdr_cells[1].text = "Price"
    
    total_price = 0.0
    for item in data.get("pricingItems", []):
        row_cells = pricing_table.add_row().cells
        row_cells[0].text = str(item.get("description", ""))
        price_str = str(item.get("price", "0"))
        row_cells[1].text = price_str
        
        try:
            price_val = float(price_str.replace("$", "").replace(",", ""))
            total_price += price_val
        except:
            pass
    
    # Total row
    row_cells = pricing_table.add_row().cells
    row_cells[0].text = "TOTAL"
    row_cells[1].text = f"${total_price:,.2f}"
    
    doc.add_page_break()
    
    # PAGE 3: MACHINES & BOM
    doc.add_heading("3. Product Scope of Supply", level=1)
    doc.add_heading("3.1 Parameter List", level=2)
    doc.add_paragraph("This proposal is based upon assumptions in Table 2.")
    
    # Machines table
    machines_table = doc.add_table(rows=1, cols=5)
    machines_table.style = 'Light Grid Accent 1'
    hdr_cells = machines_table.rows[0].cells
    hdr_cells[0].text = "Machine"
    hdr_cells[1].text = "Parameter Description"
    hdr_cells[2].text = "Parameter Quantity"
    hdr_cells[3].text = "UMM Monitors"
    hdr_cells[4].text = "VC-8000 MPS"
    
    machines = data.get("machines", [])
    for machine in machines:
        for param in machine.get("parameters", []):
            row_cells = machines_table.add_row().cells
            row_cells[0].text = str(machine.get("name", ""))
            row_cells[1].text = str(param.get("type", ""))
            row_cells[2].text = str(param.get("quantity", ""))
            row_cells[3].text = str(param.get("monitors", ""))
            row_cells[4].text = str(param.get("mps", ""))
    
    doc.add_paragraph()
    doc.add_heading("3.2 General Bill of Material", level=2)
    doc.add_paragraph("Vendor will provide the following parts.")
    
    # BOM by group
    bom_items = data.get("bomItems", [])
    groups = {}
    for item in bom_items:
        group = item.get("group", "General")
        if group not in groups:
            groups[group] = []
        groups[group].append(item)
    
    for group_name, items in groups.items():
        doc.add_paragraph(f"{group_name} Group").runs[0].bold = True
        
        bom_table = doc.add_table(rows=1, cols=4)
        bom_table.style = 'Light Grid Accent 1'
        hdr_cells = bom_table.rows[0].cells
        hdr_cells[0].text = "Item"
        hdr_cells[1].text = "Description"
        hdr_cells[2].text = "Part Number"
        hdr_cells[3].text = "Qty"
        
        for item in items:
            row_cells = bom_table.add_row().cells
            row_cells[0].text = str(item.get("item", ""))
            row_cells[1].text = str(item.get("description", ""))
            row_cells[2].text = str(item.get("partNumber", ""))
            row_cells[3].text = str(item.get("quantity", ""))
    
    doc.add_page_break()
    
    # PAGE 4: ASSUMPTIONS & SERVICES
    doc.add_heading("3.3 Product Assumptions", level=2)
    doc.add_paragraph("Please note the responsibilities table for necessary PI licensing, hardware and software.")
    
    assumptions_table = doc.add_table(rows=1, cols=3)
    assumptions_table.style = 'Light Grid Accent 1'
    hdr_cells = assumptions_table.rows[0].cells
    hdr_cells[0].text = "Scope Description"
    hdr_cells[1].text = "Buyer"
    hdr_cells[2].text = "Vendor"
    
    for item in data.get("productAssumptions", []):
        row_cells = assumptions_table.add_row().cells
        row_cells[0].text = str(item.get("description", ""))
        row_cells[1].text = "X" if item.get("buyer") else ""
        row_cells[2].text = "X" if item.get("vendor") else ""
    
    doc.add_paragraph()
    doc.add_heading("4.1 Services Responsibilities", level=2)
    doc.add_paragraph("This proposal is based upon the following responsibility table.")
    
    services_table = doc.add_table(rows=1, cols=5)
    services_table.style = 'Light Grid Accent 1'
    hdr_cells = services_table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "N/A"
    hdr_cells[3].text = "Vendor"
    hdr_cells[4].text = "Buyer"
    
    for item in data.get("servicesResponsibilities", []):
        row_cells = services_table.add_row().cells
        row_cells[0].text = str(item.get("item", ""))
        row_cells[1].text = str(item.get("description", ""))
        row_cells[2].text = "X" if item.get("na") else ""
        row_cells[3].text = "X" if item.get("vendor") else ""
        row_cells[4].text = "X" if item.get("buyer") else ""
    
    doc.add_page_break()
    
    # PAGE 5: PRODUCT DESCRIPTIONS
    doc.add_heading("5. Product Descriptions", level=1)
    
    doc.add_heading("5.1 VC-8000 Machinery Protection System", level=2)
    doc.add_paragraph(
        "The VC-8000 System is a rack-based continuous machinery monitoring platform designed to fully "
        "comply with API 670 requirements for machinery protection systems."
    )
    
    doc.add_heading("5.2 SETPOINT Condition Monitoring Software", level=2)
    doc.add_paragraph(
        "SETPOINT CMS provides collection, storage, and visualization of vibration and condition data."
    )
    
    doc.add_page_break()
    
    # PAGE 6: TERMS & EXCEPTIONS
    doc.add_heading("6. Proposal Terms", level=1)
    doc.add_paragraph("Proposal Validity: 60 days from proposal date")
    doc.add_paragraph("Payment: Net 30 days from invoice date")
    doc.add_paragraph("Limited Warranty: 36 months from invoice date")
    
    doc.add_heading("7. Exceptions and Clarifications", level=1)
    exceptions_text = data.get("exceptions", "None at this time.")
    doc.add_paragraph(exceptions_text if exceptions_text else "None at this time.")
    
    doc.add_paragraph("Thank you for considering Brüel & Kjær Vibro for this project.")
    doc.add_paragraph(data.get("salesName", "Sales Rep"))
    
    return doc

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"}), 200

@app.route("/generate-proposal", methods=["POST"])
def generate_proposal():
    try:
        data = request.get_json()
        doc = create_proposal_doc(data)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"Proposal_{data.get('proposalNumber', 'BK-Vibro')}.docx"
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
