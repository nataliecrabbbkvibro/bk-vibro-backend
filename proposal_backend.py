from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from io import BytesIO
import os
import copy

app = Flask(__name__)
CORS(app)

# Change this only if your actual template filename is different in GitHub
TEMPLATE_PATH = "Proposal_Template__2019-09-25__CODED.docx"


def normalize_text(value):
    return " ".join(str(value).strip().lower().split())


def iter_all_paragraphs(doc):
    # Main document paragraphs
    for p in doc.paragraphs:
        yield p

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

    # Headers and footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def replace_text_in_paragraph(paragraph, replacements):
    """
    Replaces text at paragraph level and rebuilds runs simply.
    This is not perfect for very complex formatting, but works well
    for your coded placeholders and footer text.
    """
    original_text = paragraph.text
    new_text = original_text

    for key, value in replacements.items():
        if key in new_text:
            new_text = new_text.replace(key, str(value))

    if new_text != original_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = new_text


def replace_everywhere(doc, replacements):
    for paragraph in iter_all_paragraphs(doc):
        replace_text_in_paragraph(paragraph, replacements)


def table_text(table):
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def find_table_containing(doc, marker):
    marker_norm = normalize_text(marker)
    for table in doc.tables:
        if marker_norm in normalize_text(table_text(table)):
            return table
    return None


def find_table_with_any_marker(doc, markers):
    for marker in markers:
        table = find_table_containing(doc, marker)
        if table is not None:
            return table
    return None


def clear_table_rows(table, keep_rows=1):
    while len(table.rows) > keep_rows:
        tr = table.rows[-1]._element
        tr.getparent().remove(tr)


def set_cell_text(cell, value):
    cell.text = str(value)


def safe_float(value):
    try:
        return float(str(value).replace("$", "").replace(",", "").strip())
    except Exception:
        return 0.0


def fill_revision_history_table(doc, revision_items):
    table = find_table_with_any_marker(doc, ["<REV_DATE>", "Revision", "Description"])
    if table is None:
        return

    # Keep header row only, then rebuild
    clear_table_rows(table, keep_rows=1)

    if not revision_items:
        revision_items = [{"date": "", "revision": "", "description": ""}]

    for item in revision_items:
        row = table.add_row()
        set_cell_text(row.cells[0], item.get("date", ""))
        set_cell_text(row.cells[1], item.get("revision", ""))
        set_cell_text(row.cells[2], item.get("description", ""))


def fill_pricing_table(doc, pricing_items):
    table = find_table_with_any_marker(doc, ["<PRICE_SCOPE>", "Scope of Supply Description"])
    if table is None:
        return 0.0

    clear_table_rows(table, keep_rows=1)

    total = 0.0
    for item in pricing_items:
        row = table.add_row()
        desc = item.get("description", "")
        price = item.get("price", "")
        set_cell_text(row.cells[0], desc)
        set_cell_text(row.cells[1], price)
        total += safe_float(price)

    return total


def fill_machines_table(doc, machines):
    table = find_table_with_any_marker(doc, ["<M1_NAME>", "Parameter Description", "Parameter Quantity"])
    if table is None:
        return

    clear_table_rows(table, keep_rows=1)

    total_qty = 0.0
    total_umm = 0.0
    total_mps = 0.0

    for machine in machines:
        name = machine.get("name", "")
        params = machine.get("parameters", [])
        first_row = True

        for param in params:
            row = table.add_row()
            set_cell_text(row.cells[0], name if first_row else "")
            set_cell_text(row.cells[1], param.get("type", ""))
            set_cell_text(row.cells[2], param.get("quantity", ""))
            set_cell_text(row.cells[3], param.get("monitors", ""))
            set_cell_text(row.cells[4], param.get("mps", ""))

            total_qty += safe_float(param.get("quantity", 0))
            total_umm += safe_float(param.get("monitors", 0))
            total_mps += safe_float(param.get("mps", 0))

            first_row = False

    # Totals row
    row = table.add_row()
    set_cell_text(row.cells[0], "Totals:")
    set_cell_text(row.cells[1], "")
    set_cell_text(row.cells[2], int(total_qty) if total_qty.is_integer() else total_qty)
    set_cell_text(row.cells[3], int(total_umm) if total_umm.is_integer() else total_umm)
    set_cell_text(row.cells[4], int(total_mps) if total_mps.is_integer() else total_mps)


def fill_bom_table(doc, bom_items):
    table = find_table_with_any_marker(doc, ["<MPS_I1>", "Part Number", "Quantity"])
    if table is None:
        return

    clear_table_rows(table, keep_rows=1)

    for item in bom_items:
        row = table.add_row()
        set_cell_text(row.cells[0], item.get("group", ""))
        set_cell_text(row.cells[1], item.get("item", ""))
        set_cell_text(row.cells[2], item.get("description", ""))
        set_cell_text(row.cells[3], item.get("partNumber", ""))
        set_cell_text(row.cells[4], item.get("quantity", ""))


def update_product_assumptions_table(doc, product_assumptions):
    """
    Keep template descriptions exactly as written.
    Only change X in Buyer/Vendor columns.
    """
    table = find_table_with_any_marker(
        doc,
        ["Local PI Server Hardware", "PI System Access (PSA) for SETPOINT"]
    )
    if table is None:
        return

    incoming = {
        normalize_text(item.get("description", "")): item
        for item in product_assumptions
    }

    # Table columns expected:
    # 0 = Buyer (or Others)
    # 1 = Vendor
    # 2 = Scope Description
    for row in table.rows[1:]:
        desc = normalize_text(row.cells[2].text)
        item = incoming.get(desc)
        if item:
            set_cell_text(row.cells[0], "X" if item.get("buyer") else "")
            set_cell_text(row.cells[1], "X" if item.get("vendor") else "")
            # Keep row.cells[2] unchanged on purpose


def fill_scope_services_table(doc, scope_services):
    """
    Table 4 - Services Scope of Supply
    Expandable.
    """
    table = find_table_with_any_marker(doc, ["<SVC_ITEM1>", "Vendor will provide the following services"])
    if table is None:
        # fallback by matching the coded placeholder in any table cell
        table = find_table_with_any_marker(doc, ["<SVC_DESC1>", "<SVC_QTY1>"])
    if table is None:
        return

    clear_table_rows(table, keep_rows=1)

    if not scope_services:
        scope_services = [{"group": "Services", "item": "", "description": "", "quantity": ""}]

    for item in scope_services:
        row = table.add_row()
        set_cell_text(row.cells[0], item.get("group", "Services"))
        set_cell_text(row.cells[1], item.get("item", ""))
        set_cell_text(row.cells[2], item.get("description", ""))
        set_cell_text(row.cells[3], item.get("quantity", ""))


def update_services_responsibilities_table(doc, responsibilities):
    """
    Keep template wording exactly as written.
    Only change X values.
    """
    table = find_table_with_any_marker(
        doc,
        ["Receiving, off-loading and storage of Vendor provided parts", "Machine disassembly/assembly"]
    )
    if table is None:
        return

    incoming = {
        normalize_text(item.get("description", "")): item
        for item in responsibilities
    }

    # Expected columns:
    # 0 = Item
    # 1 = Description
    # 2 = N/A
    # 3 = Vendor
    # 4 = Buyer
    for row in table.rows[1:]:
        desc = normalize_text(row.cells[1].text)
        item = incoming.get(desc)

        if item:
            na = bool(item.get("na"))
            vendor = bool(item.get("vendor"))
            buyer = bool(item.get("buyer"))

            # Keep one X choice clean
            if na:
                vendor = False
                buyer = False
            elif vendor:
                na = False
                buyer = False
            elif buyer:
                na = False
                vendor = False

            set_cell_text(row.cells[2], "X" if na else "")
            set_cell_text(row.cells[3], "X" if vendor else "")
            set_cell_text(row.cells[4], "X" if buyer else "")


def fill_exceptions_table(doc, exceptions):
    table = find_table_with_any_marker(doc, ["Existing Wording", "Suggested Wording", "Notes"])
    if table is None:
        return

    clear_table_rows(table, keep_rows=1)

    if not exceptions:
        return

    for item in exceptions:
        row = table.add_row()
        set_cell_text(row.cells[0], item.get("section", ""))
        set_cell_text(row.cells[1], item.get("existingWording", ""))
        set_cell_text(row.cells[2], item.get("suggestedWording", ""))
        set_cell_text(row.cells[3], item.get("notes", ""))


def generate_proposal(data):
    if not os.path.exists(TEMPLATE_PATH):
        raise Exception(f"Template not found: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)

    pricing_items = data.get("pricingItems", [])
    revision_history = data.get("revisionHistory", [])
    machines = data.get("machines", [])
    bom_items = data.get("bomItems", [])
    product_assumptions = data.get("productAssumptions", [])
    scope_services = data.get("scopeServices", [])
    services_responsibilities = data.get("servicesResponsibilities", [])
    exceptions = data.get("exceptions", [])

    total_price = fill_pricing_table(doc, pricing_items)

    # Cover page, revision placeholders, footer labels, etc.
    replacements = {
        "<BUYER_CONTACT_NAME>": data.get("buyerContactName", ""),
        "<BUYER_CONTACT_TITLE>": data.get("buyerContactTitle", ""),
        "<BUYER_COMPANY>": data.get("buyerCompanyName", ""),
        "<BUYER_ADDRESS>": data.get("buyerCompanyAddress", ""),
        "<BUYER_PHONE>": data.get("buyerContactPhone", ""),
        "<BUYER_EMAIL>": data.get("buyerContactEmail", ""),
        "<SALES_NAME>": data.get("salesName", ""),
        "<SALES_TITLE>": data.get("salesTitle", ""),
        "<SALES_PHONE>": data.get("salesPhone", ""),
        "<SALES_EMAIL>": data.get("salesEmail", ""),

        # Revision placeholders in case any remain elsewhere
        "<REV_DATE>": revision_history[0].get("date", "") if revision_history else "",
        "<REV_NO>": revision_history[0].get("revision", "") if revision_history else "",
        "<REV_DESC>": revision_history[0].get("description", "") if revision_history else "",

        # Pricing placeholders if any remain elsewhere
        "<PRICE_SCOPE>": "; ".join(
            str(item.get("description", "")) for item in pricing_items if item.get("description")
        ),
        "<TOTAL_PRICE>": f"${total_price:,.2f}" if total_price else "",

        # Footer / legacy labels
        "Proposal Number": data.get("proposalNumber", ""),
        "Buyer Name": data.get("buyerCompanyName", ""),
        "Project Name": data.get("projectName", ""),
        "Proposal Date": data.get("proposalDate", ""),
    }

    replace_everywhere(doc, replacements)

    fill_revision_history_table(doc, revision_history)
    fill_machines_table(doc, machines)
    fill_bom_table(doc, bom_items)
    update_product_assumptions_table(doc, product_assumptions)
    fill_scope_services_table(doc, scope_services)
    update_services_responsibilities_table(doc, services_responsibilities)
    fill_exceptions_table(doc, exceptions)

    return doc


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"}), 200


@app.route("/generate-proposal", methods=["POST"])
def generate_proposal_route():
    try:
        data = request.get_json()
        doc = generate_proposal(data)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"Proposal_{data.get('proposalNumber', 'BK-Vibro')}.docx"
        )
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
